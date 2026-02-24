"""
Module de génération de documents pour le projet Publipostage DOETH.

Ce module gère la création des attestations DOETH au format Word et/ou PDF
à partir des données CSV prétraitées.
"""
import csv
import logging
import os
import tempfile
import datetime
from enum import Enum
from pathlib import Path
from typing import Optional, List, Dict, Any

import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT as WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT as WD_ALIGN_VERTICAL
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from src.utils.config import get
from src.utils.logger import get_logger

# Initialisation du logger
logger = get_logger(__name__)


class OutputFormat(Enum):
    """Format de sortie des attestations."""
    DOCX = "docx"
    PDF = "pdf"
    BOTH = "both"


def create_document(template_path: Optional[str] = None) -> Document:
    """
    Crée un nouveau document Word avec les paramètres par défaut.

    Args:
        template_path: Chemin vers un modèle de document (optionnel)

    Returns:
        Document: Document Word initialisé
    """
    logger.debug(f"Création d'un nouveau document Word")

    if template_path and os.path.exists(template_path):
        logger.debug(f"Utilisation du modèle: {template_path}")
        doc = Document(template_path)
    else:
        doc = Document()

    # Configurer les marges du document
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(get('document.margins.top', 1.0))
        section.bottom_margin = Cm(get('document.margins.bottom', 1.0))
        section.left_margin = Cm(get('document.margins.left', 1.5))
        section.right_margin = Cm(get('document.margins.right', 1.5))

    # Définir la taille de police par défaut
    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(get('document.font_size', 10))

    # Réduire l'espacement après les paragraphes
    style.paragraph_format.space_after = Pt(
        get('document.paragraph_spacing', 4))

    return doc


def add_logo(doc: Document, logo_path: Optional[str] = None) -> None:
    """
    Ajoute un logo en en-tête du document.

    Args:
        doc: Document Word à modifier
        logo_path: Chemin vers l'image du logo
    """
    if not logo_path or not os.path.exists(logo_path):
        logo_path = get('resources.logo_path')
        if not logo_path or not os.path.exists(logo_path):
            logger.warning("Logo non trouvé, en-tête non ajouté")
            return

    logger.debug(f"Ajout du logo: {logo_path}")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    logo_width = Cm(get('document.logo_width', 4.0))
    run.add_picture(logo_path, width=logo_width)


def add_client_header(doc: Document, client_info: Dict[str, Any]) -> None:
    """
    Ajoute les informations du client en en-tête à droite.

    Args:
        doc: Document Word à modifier
        client_info: Dictionnaire contenant les informations du client
    """
    logger.debug("Ajout des informations client en en-tête")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Nom client
    nom_client = client_info.get('NOM_CLIENT', '')
    if nom_client:
        run = p.add_run(f"{nom_client}")
        run.bold = True
        p.add_run("\n")

    # Adresse
    adresse = client_info.get('ADRESSE CLIENT', '')
    if adresse:
        run = p.add_run(f"{adresse}")
        run.bold = True
        p.add_run("\n")

    # Code postal
    cp = client_info.get('CP CLIENT', '')
    if cp:
        run = p.add_run(f"{cp}")
        run.bold = True
        p.add_run("\n")

    # Ville
    ville = client_info.get('VILLE CLIENT', '')
    if ville:
        run = p.add_run(f"{ville}")
        run.bold = True


def add_title(doc: Document) -> None:
    """
    Ajoute le titre centré au document.

    Args:
        doc: Document Word à modifier
    """
    logger.debug("Ajout du titre")

    # Ajouter un espace avant le titre
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    title_text = get('document.title',
                     "Attestation relative aux travailleurs en situation d'handicap mis à disposition "
                     "par une entreprise de travail temporaire ou un groupement d'employeurs"
                     )

    title_run = p.add_run(title_text)
    title_run.bold = True


def add_empty_space(doc: Document, size: int = 4) -> None:
    """
    Ajoute un espace (paragraphe vide avec police réduite) pour gérer l'espacement.

    Args:
        doc: Document Word à modifier
        size: Taille de police du paragraphe vide (plus petite = espace réduit)
    """
    p_empty = doc.add_paragraph()
    run_empty = p_empty.add_run()
    run_empty.font.size = Pt(size)


def add_legal_references(doc: Document) -> None:
    """
    Ajoute les références légales au document.

    Args:
        doc: Document Word à modifier
    """
    logger.debug("Ajout des références légales")

    legal_text = get('document.legal_reference',
                     "Vu les articles L5212-1, D5212-1, D5212-3, D5212-6 et D5212-8 du Code du travail,"
                     )

    # Ajouter un espace avant le titre
    doc.add_paragraph()

    p = doc.add_paragraph(legal_text)


def add_representative_info(doc: Document) -> None:
    """
    Ajoute les informations du représentant légal.

    Args:
        doc: Document Word à modifier
    """
    logger.debug("Ajout des informations du représentant légal")

    doc.add_paragraph()

    rep_name = get('representant.nom', "Loïc GALLERAND")
    rep_adresse = get('representant.adresse',
                      "233 rue de Châteaugiron à Rennes (35000)")
    rep_siret = get('representant.siret', "49342093900057")

    p = doc.add_paragraph(f"Je soussigné, {rep_name}")
    p.add_run("\nReprésentant légal de l'entreprise de travail temporaire située au")
    p.add_run(f"\n{rep_adresse}")
    p.add_run(f"\nSIRET : {rep_siret}")


def add_client_attestation(doc: Document, client_info: Dict[str, Any]) -> None:
    """
    Ajoute l'attestation avec les informations du client.

    Args:
        doc: Document Word à modifier
        client_info: Dictionnaire contenant les informations du client
    """
    logger.debug("Ajout de l'attestation client")

    add_empty_space(doc)

    # Phrase "Atteste que"
    p = doc.add_paragraph("Atteste que")

    # Informations du client
    p = doc.add_paragraph()
    nom_client = client_info.get('NOM_CLIENT', '')
    siret = client_info.get('SIRET', '')

    p.add_run(f"Nom client : {nom_client}").bold = True
    p.add_run("\n")
    p.add_run(f"SIRET : {siret}").bold = True

    # Année -1 dans le texte de l'attestation
    year = datetime.datetime.now().year - 1

    # Paragraphe explicatif
    doc.add_paragraph()
    explanation_text = get(f'document.explanation_text_{year}',
                           f"Peut, valoriser, dans le cadre de la déclaration obligatoire d'emploi des travailleurs "
                           f"en situation d'handicap au titre de l'année civile {year} les bénéficiaires de l'obligation "
                           f"d'emploi des travailleurs handicapés mis à disposition suivants :"
                           )

    p = doc.add_paragraph(explanation_text)


def create_employee_table(doc: Document, employees_data: pd.DataFrame) -> None:
    """
    Crée et remplit le tableau des employés.

    Args:
        doc: Document Word à modifier
        employees_data: DataFrame contenant les données des employés
    """
    logger.debug(f"Création du tableau pour {len(employees_data)} employés")

    doc.add_paragraph()
    table = doc.add_table(rows=1, cols=8)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # En-têtes du tableau
    headers = ["REGROUPEMENT", "SIRET", "PRENOM", "NOM", "QUALIFICATION",
               "ETP_MAJORE", "Nombre d'heure", "ETP annuelle"]

    # Appliquer le style pour les en-têtes
    hdr_cells = table.rows[0].cells
    table_font_size = get('document.table_font_size', 8)

    for j, header in enumerate(headers):
        hdr_cells[j].text = header
        run = hdr_cells[j].paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(table_font_size)

        # Ajouter un fond gris clair
        shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D3D3D3"/>')
        hdr_cells[j]._tc.get_or_add_tcPr().append(shading_elm)

        # Centrer verticalement et horizontalement
        hdr_cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hdr_cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Ajout des lignes pour chaque employé
    total_etp = 0
    total_heures = 0

    for _, row in employees_data.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(
            row['REGROUPEMENT']) if row['REGROUPEMENT'] == row['REGROUPEMENT'] else ''
        row_cells[1].text = str(row['SIRET'])
        row_cells[2].text = str(
            row['PRENOM']) if row['PRENOM'] == row['PRENOM'] else ''
        row_cells[3].text = str(row['NOM']) if row['NOM'] == row['NOM'] else ''
        row_cells[4].text = str(
            row['QUALIFICATION']) if row['QUALIFICATION'] == row['QUALIFICATION'] else ''
        row_cells[5].text = str(
            row['ETP_MAJORE']) if row['ETP_MAJORE'] == row['ETP_MAJORE'] else ''
        row_cells[6].text = str(
            row['NB_HEURES']) if row['NB_HEURES'] == row['NB_HEURES'] else ''
        row_cells[7].text = f"{float(row['ETP_ANNUEL']):.2f}" if row['ETP_ANNUEL'] == row['ETP_ANNUEL'] else ''

        # Appliquer la taille de police 8pt à toutes les cellules de données
        for col in range(8):
            for paragraph in row_cells[col].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(table_font_size)

        # Centrer certaines cellules
        for col in [6, 7]:  # ETP_MAJORE, NB_HEURES, ETP_ANNUEL
            row_cells[col].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # Ajouter au total
        total_etp += float(row['ETP_ANNUEL']
                           ) if row['ETP_ANNUEL'] == row['ETP_ANNUEL'] else 0.0
        total_heures += float(row['NB_HEURES']
                              ) if row['NB_HEURES'] == row['NB_HEURES'] else 0.0

    # Ajouter la ligne de total
    row_cells = table.add_row().cells
    row_cells[0].merge(row_cells[6])  # Fusionner les 7 premières cellules
    row_cells[0].text = "Total d'unités bénéficiaires"
    row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    # Afficher le total avec 2 décimales
    row_cells[7].text = f"{total_etp:.2f}"
    row_cells[7].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Appliquer la taille de police à la ligne de total
    for run in row_cells[0].paragraphs[0].runs:
        run.font.size = Pt(table_font_size)
    for run in row_cells[7].paragraphs[0].runs:
        run.font.size = Pt(table_font_size)

    # Ajuster la largeur des colonnes
    column_widths = [3.0, 2.0, 2.2, 2.5, 2.5, 1.5, 1.8, 1.8]  # en cm

    for col, width in enumerate(column_widths):
        if col < len(table.columns):
            for cell in table.columns[col].cells:
                cell.width = Cm(width)


def add_footer_and_signature(doc: Document, signature_path: Optional[str] = None) -> None:
    """
    Ajoute le pied de page avec la date, le représentant légal et la signature.

    Args:
        doc: Document Word à modifier
        signature_path: Chemin vers l'image de signature
    """
    logger.debug("Ajout du pied de page et de la signature")

    add_empty_space(doc)

    # Date actuelle
    date_format = get('defaults.date_format', "%d/%m/%Y")
    date_now = datetime.datetime.now().strftime(date_format)
    city = get('document.city', "Rennes")

    p = doc.add_paragraph(f"Fait à {city}, le {date_now}")

    # Représentant légal
    doc.add_paragraph()
    p = doc.add_paragraph("Le représentant légal,")

    # Espace pour la signature
    add_empty_space(doc)

    # Ajouter la signature si fournie
    if not signature_path or not os.path.exists(signature_path):
        signature_path = get('resources.signature_path')
        if not signature_path or not os.path.exists(signature_path):
            logger.warning("Image de signature non trouvée")
            return

    logger.debug(f"Ajout de la signature: {signature_path}")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    signature_width = Cm(get('document.signature_width', 4.5))
    run.add_picture(signature_path, width=signature_width)


def save_document(doc: Document, output_path: str) -> str:
    """
    Enregistre le document Word.

    Args:
        doc: Document Word à enregistrer
        output_path: Chemin où enregistrer le document

    Returns:
        str: Chemin du document enregistré
    """
    # Créer le répertoire de sortie si nécessaire
    output_dir = os.path.dirname(output_path)
    if output_dir and not os.path.exists(output_dir):
        os.makedirs(output_dir, exist_ok=True)

    logger.debug(f"Sauvegarde du document: {output_path}")
    try:
        doc.save(output_path)
        logger.info(f"Document enregistré: {output_path}")
        return output_path
    except Exception as e:
        logger.error(f"Erreur lors de l'enregistrement du document: {str(e)}")
        raise


def create_attestation(file_number: int, siret_data: pd.DataFrame, output_folder: str,
                       logger: logging.Logger,
                       signature_path: Optional[str] = None,
                       logo_path: Optional[str] = None,
                       output_format: OutputFormat = OutputFormat.DOCX) -> List[str]:
    """
    Crée une attestation DOETH pour un SIRET donné.

    Le DOCX est toujours généré en premier (source de vérité du layout).
    Si output_format est PDF, le DOCX est converti puis supprimé.
    Si output_format est BOTH, les deux fichiers sont conservés.

    Args:
        file_number: N° de fichier à créer
        siret_data: DataFrame contenant les données d'un SIRET
        output_folder: Dossier où enregistrer l'attestation
        logger: Logger pour suivre l'évolution du processus
        signature_path: Chemin vers l'image de signature
        logo_path: Chemin vers l'image du logo
        output_format: Format de sortie souhaité (DOCX, PDF ou BOTH)

    Returns:
        List[str]: Liste des chemins des fichiers générés
    """
    if len(siret_data) == 0:
        logger.warning(
            "Aucune donnée fournie pour la génération de l'attestation")
        return []

    # Récupérer les informations générales (première ligne)
    info = siret_data.iloc[0]
    siret = info['SIRET']
    nom_client = info['NOM_CLIENT']
    nom_regroupement = info['REGROUPEMENT']

    logger.info(f"Création de l'attestation pour SIRET: {siret}, Client: {nom_client}, "
                f"Nom du regroupement: {nom_regroupement}")

    # Construction du document Word
    doc = create_document()
    add_logo(doc, logo_path)
    add_client_header(doc, info)
    add_title(doc)
    add_legal_references(doc)
    add_representative_info(doc)
    add_client_attestation(doc, info)
    create_employee_table(doc, siret_data)
    add_footer_and_signature(doc, signature_path)

    # Nom de base du fichier (sans extension)
    year = datetime.datetime.now().year - 1
    base_name = f"{file_number}_Attestation DOETH_{year}_{nom_regroupement}"

    # Sauvegarde du DOCX (toujours nécessaire comme intermédiaire)
    docx_path = Path(output_folder) / f"{base_name}.docx"
    save_document(doc, str(docx_path))

    generated: List[str] = []

    if output_format in (OutputFormat.DOCX, OutputFormat.BOTH):
        generated.append(str(docx_path))

    if output_format in (OutputFormat.PDF, OutputFormat.BOTH):
        pdf_path = _convert_to_pdf(docx_path, logger)
        if pdf_path:
            generated.append(str(pdf_path))
        # Supprimer le DOCX intermédiaire si seul le PDF est demandé
        if output_format == OutputFormat.PDF:
            try:
                docx_path.unlink()
            except OSError as e:
                logger.warning(
                    f"Impossible de supprimer le DOCX intermédiaire: {e}")

    return generated


def _convert_to_pdf(docx_path: Path, logger: logging.Logger) -> Optional[Path]:
    """
    Convertit un fichier DOCX en PDF via Word COM (Windows) ou LibreOffice (Linux/macOS).

    Args:
        docx_path: Chemin du fichier DOCX source
        logger: Logger

    Returns:
        Path du PDF généré, ou None en cas d'échec
    """
    pdf_path = docx_path.with_suffix('.pdf')
    try:
        from docx2pdf import convert
        convert(str(docx_path), str(pdf_path))
        logger.debug(f"PDF généré: {pdf_path}")
        return pdf_path
    except ImportError:
        logger.error(
            "docx2pdf n'est pas installé. Installez-le avec : uv add docx2pdf")
        return None
    except Exception as e:
        logger.error(
            f"Erreur lors de la conversion PDF pour {docx_path.name}: {e}")
        return None


def generer_attestations_doeth(csv_path: str, output_folder: str,
                               logger: logging.Logger,
                               signature_path: Optional[str] = None,
                               logo_path: Optional[str] = None,
                               output_format: OutputFormat = OutputFormat.DOCX) -> List[str]:
    """
    Génère des attestations DOETH regroupées par SIRET à partir d'un fichier CSV.

    Args:
        csv_path: Chemin vers le fichier CSV source
        output_folder: Dossier où enregistrer les attestations
        logger: Logger pour suivre l'évolution des opérations
        signature_path: Chemin vers l'image de signature
        logo_path: Chemin vers l'image du logo
        output_format: Format de sortie (DOCX, PDF ou BOTH)

    Returns:
        List[str]: Liste des chemins des fichiers générés
    """
    format_label = {
        OutputFormat.DOCX: "Word (.docx)",
        OutputFormat.PDF: "PDF (.pdf)",
        OutputFormat.BOTH: "Word + PDF",
    }[output_format]

    logger.info(
        f"Génération des attestations ({format_label}) depuis: {csv_path}")
    os.makedirs(output_folder, exist_ok=True)

    try:
        separator = get('defaults.csv_separator', ';')
        df = pd.read_csv(csv_path, sep=separator, quoting=csv.QUOTE_NONNUMERIC,
                         dtype={'SIRET': str, 'SIREN': str, 'NIC': str})
        logger.info(f"Fichier CSV chargé: {len(df)} lignes")

        df = df.sort_values(by=['SIRET', 'NOM', 'PRENOM'])
        sirets = df['SIRET'].unique()
        logger.info(f"Nombre total de SIRET à traiter: {len(sirets)}")

        generated_docs: List[str] = []

        for i, siret in enumerate(sirets):
            siret_data = df[df['SIRET'] == siret]
            if len(siret_data) == 0:
                logger.warning(f"Aucune donnée pour le SIRET: {siret}")
                continue

            files = create_attestation(
                i + 1, siret_data, output_folder, logger,
                signature_path, logo_path, output_format)
            generated_docs.extend(files)

            # Barre de progression
            progress = (i + 1) / len(sirets)
            filled = int(30 * progress)
            bar = '█' * filled + '░' * (30 - filled)
            logger.info(
                f"({i + 1}/{len(sirets)}) [{bar}] {progress:.1%}: {', '.join(Path(f).name for f in files)}")

        logger.info(
            f"Génération terminée: {len(generated_docs)} fichier(s) dans {output_folder}")
        return generated_docs

    except Exception as e:
        logger.error(f"Erreur lors de la génération des attestations: {e}")
        raise
